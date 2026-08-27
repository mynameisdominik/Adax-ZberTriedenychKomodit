from __future__ import annotations

import io
import re
from dataclasses import asdict, dataclass
from datetime import date
from typing import BinaryIO

import requests
from docx import Document

DATE_PATTERN = re.compile(r"\b(\d{1,2})\.(\d{1,2})\.(\d{4})\b")
PLASTIC_NAMES = ("plast", "vkm", "kov")
PAPER_NAME = "papier"


@dataclass(frozen=True)
class Collection:
    date: date
    commodity: str

    def as_dict(self) -> dict[str, str]:
        result = asdict(self)
        result["date"] = self.date.isoformat()
        return result


def parse_url(url: str, *, timeout: int = 30) -> list[Collection]:
    response = requests.get(url, timeout=timeout)
    response.raise_for_status()
    return parse_docx(io.BytesIO(response.content))


def parse_docx(source: str | bytes | BinaryIO) -> list[Collection]:
    if isinstance(source, str):
        document = Document(source)
    elif isinstance(source, bytes):
        document = Document(io.BytesIO(source))
    else:
        document = Document(source)

    collections: list[Collection] = []
    for table in document.tables:
        if not table.rows:
            continue
        headers = [_normalize(cell.text) for cell in table.rows[0].cells]
        commodity_columns = {
            index: _commodity(header)
            for index, header in enumerate(headers)
            if _commodity(header) is not None
        }
        if not commodity_columns:
            continue

        for row in table.rows[1:]:
            for index, commodity in commodity_columns.items():
                if index >= len(row.cells):
                    continue
                for match in DATE_PATTERN.finditer(row.cells[index].text):
                    collections.append(
                        Collection(
                            date=date(
                                int(match.group(3)),
                                int(match.group(2)),
                                int(match.group(1)),
                            ),
                            commodity=commodity,
                        )
                    )
        if collections:
            return collections

    raise ValueError("No table with supported ADAX commodities was found")


def _normalize(value: str) -> str:
    return " ".join(value.casefold().split())


def _commodity(header: str) -> str | None:
    if PAPER_NAME in header:
        return "Papier"
    if all(name in header for name in PLASTIC_NAMES):
        return "Plast, VKM, Kov ob."
    return None
